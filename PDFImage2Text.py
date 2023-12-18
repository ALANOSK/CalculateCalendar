from PIL import Image
import pytesseract
import os
import io
from wand.image import Image as wi
from datetime import datetime
from docx import Document

# Set the Tesseract command path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def process_image(file_path):
    img = Image.open(file_path)
    text = pytesseract.image_to_string(img, lang='eng')
    return text

def process_scanned_pdf(pdf_path):
    pdf_file = wi(filename=pdf_path, resolution=300)
    image_sequence = pdf_file.sequence
    
    extracted_text = []
    for index, img in enumerate(image_sequence):
        img_page = wi(image=img)
        img_blob = img_page.make_blob('jpeg')
        
        img = Image.open(io.BytesIO(img_blob))
        config = r'--tessdata-dir "C:\Program Files\Tesseract-OCR\tessdata"'
        text = pytesseract.image_to_string(img, lang='eng', config=config)
        extracted_text.append(text)
    
    # Close the PDF file
    pdf_file.close()
    
    return extracted_text

def save_to_word_document(extracted_text, image_path=None, pdf_path=None):
    doc = Document()

    for index, text in enumerate(extracted_text):
        doc.add_heading(f'Page {index + 1}', level=1)
        doc.add_paragraph(text)

    # Extract the original filename without extension from the provided path
    if image_path:
        original_name = os.path.splitext(os.path.basename(image_path))[0]
    elif pdf_path:
        original_name = os.path.splitext(os.path.basename(pdf_path))[0]
    else:
        original_name = "Untitled"

    # Create 'outputresult' directory if not exists
    output_dir = os.path.join(os.getcwd(), 'outputresult')
    os.makedirs(output_dir, exist_ok=True)

    # Save the document with a filename based on the original name and timestamp
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    output_filename = f"{original_name} - Image2Text ({timestamp}).docx"
    output_path = os.path.join(output_dir, output_filename)

    doc.save(output_path)
    print(f"Text saved to Word document: {output_path}")

def process_images_in_directory(directory_path):
    extracted_text = []
    for root, dirs, filenames in os.walk(directory_path):
        for filename in filenames:
            img_path = os.path.join(root, filename)
            print('#####################################', filename, '#####################################')
            text = process_image(img_path)
            extracted_text.append(text)
            print(text)
    
    return extracted_text

def main():
    print("Choose an option:")
    print("1. Process a single image")
    print("2. Process a scanned PDF")
    print("3. Process images in a directory")

    choice = input("Enter your choice (1/2/3): ")

    if choice == "1":
        image_path = input("Enter the path of the image: ")
        result = process_image(image_path)
        print(result)
    elif choice == "2":
        pdf_path = input("Enter the path of the scanned PDF: ")
        result = process_scanned_pdf(pdf_path)
        save_to_word_document(result)
    elif choice == "3":
        directory_path = input("Enter the path of the directory containing images: ")
        result = process_images_in_directory(directory_path)
        save_to_word_document(result)
    else:
        print("Invalid choice. Please enter 1, 2, or 3.")

if __name__ == "__main__":
    main()